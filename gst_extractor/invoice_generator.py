from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


@dataclass
class InvoiceDraft:
    invoice_no: str
    invoice_date: str
    party_name: str
    party_gstin: str
    party_address: str
    hsn: str = "998346"
    description: str = "FABRIC TESTING / ANALYSIS CHARGES"
    taxable: float = 0.0
    freight: float = 0.0
    cgst: float = 0.0
    sgst: float = 0.0
    igst: float = 0.0
    total: float = 0.0
    business_name: str = "DEMO TEXTILE RESEARCH LAB"
    business_gstin: str = "00ABCDE1234A1Z0"
    business_address: str = "3/5, Rajasthan Housing Board, 4th Phase RIICO Sample City-311001(Raj.)"
    email: str = "demobhilwara@gmail.com"
    mobile: str = "99999-99999"


def _money(v: Optional[float]) -> float:
    try:
        return round(float(v or 0), 2)
    except Exception:
        return 0.0


def calculate_taxes(taxable: float, freight: float = 0.0, customer_gstin: str = "", business_gstin: str = "00ABCDE1234A1Z0", special_rate: Optional[float] = None) -> tuple[float, float, float, float]:
    base = _money(taxable) + _money(freight)
    rate = 18.0 if special_rate is None else float(special_rate)
    seller_state = (business_gstin or "")[:2]
    buyer_state = (customer_gstin or "")[:2]
    if buyer_state and seller_state and buyer_state != seller_state:
        igst = round(base * rate / 100, 2)
        return round(base, 2), 0.0, 0.0, igst
    half = round(base * (rate / 2) / 100, 2)
    return round(base, 2), half, half, 0.0


def generate_invoice_docx(draft: InvoiceDraft, output_path: str | Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    taxable_total = _money(draft.taxable) + _money(draft.freight)
    total = taxable_total + _money(draft.cgst) + _money(draft.sgst) + _money(draft.igst)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TAX INVOICE")
    run.bold = True
    run.font.size = Pt(14)

    header = doc.add_table(rows=1, cols=1)
    header.style = "Table Grid"
    cell = header.cell(0, 0)
    p = cell.paragraphs[0]
    p.add_run(f"GSTIN : {draft.business_gstin}\n").bold = True
    p.add_run(f"{draft.business_name}\n").bold = True
    p.add_run(f"{draft.business_address}\n")
    p.add_run(f"Email : {draft.email} , Mobile No. : {draft.mobile}")

    details = doc.add_table(rows=6, cols=4)
    details.style = "Table Grid"
    data = [
        ("Details of Receiver (Billed to)", "", "GST Invoice No.", draft.invoice_no),
        ("Name", draft.party_name, "Invoice Date", draft.invoice_date),
        ("Address", draft.party_address, "Despatch Through", "-"),
        ("GSTIN", draft.party_gstin, "Reverse Charge(Y/N)", "-"),
        ("State", "Rajasthan" if (draft.party_gstin or "")[:2] == "08" else "", "State Code", (draft.party_gstin or "")[:2]),
        ("", "", "", ""),
    ]
    for r, row in enumerate(data):
        for c, value in enumerate(row):
            details.cell(r, c).text = str(value)
            if c in (0, 2):
                for run in details.cell(r, c).paragraphs[0].runs:
                    run.bold = True

    items = doc.add_table(rows=2, cols=6)
    items.style = "Table Grid"
    heads = ["S.N.", "Test Parameters", "HSN Code", "Quantity", "Rs./sample", "Amount(Rs.)"]
    vals = ["1", draft.description, draft.hsn, "01", f"{_money(draft.taxable):.2f}", f"{_money(draft.taxable):.2f}"]
    for c, h in enumerate(heads):
        items.cell(0, c).text = h
        for run in items.cell(0, c).paragraphs[0].runs:
            run.bold = True
    for c, v in enumerate(vals):
        items.cell(1, c).text = v

    totals = doc.add_table(rows=7, cols=2)
    totals.style = "Table Grid"
    rows = [
        ("Total Amt. Before Tax", _money(draft.taxable)),
        ("Add : Freight", _money(draft.freight)),
        ("Add: SGST", _money(draft.sgst)),
        ("Add: CGST", _money(draft.cgst)),
        ("Add: IGST", _money(draft.igst)),
        ("Total Tax Amt.", _money(draft.sgst) + _money(draft.cgst) + _money(draft.igst)),
        ("Total Amt. After Tax", _money(total)),
    ]
    for r, (label, val) in enumerate(rows):
        totals.cell(r, 0).text = label
        totals.cell(r, 1).text = "-" if val == 0 and "Total" not in label else f"{val:.2f}"
        for run in totals.cell(r, 0).paragraphs[0].runs:
            run.bold = True

    sign = doc.add_paragraph("\nCertified that the particulars given above are true and correct\n")
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph(f"For: {draft.business_name}\nAuthorised Signatory")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(output_path)
    return output_path
