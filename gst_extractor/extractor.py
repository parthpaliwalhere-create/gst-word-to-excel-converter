from __future__ import annotations

import re
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from dateutil import parser as date_parser
from docx import Document

try:
    import pdfplumber
except Exception:
    pdfplumber = None

GSTIN_RE = re.compile(r"\b[0-9]{2}[A-Z]{5}[0-9]{4}[A-Z][1-9A-Z]Z[0-9A-Z]\b", re.I)
MONEY_RE = re.compile(r"(?:₹|INR|Rs\.?\s*)?\s*([-+]?\d+(?:,\d{2,3})*(?:\.\d+)?)")
DATE_RE = re.compile(r"\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{1,2}\s+[A-Za-z]{3,9}\s+\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2})\b")
SELLER_GSTIN_DEFAULT = "00ABCDE1234A1Z0"


def clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def norm_key(value: str) -> str:
    value = clean_text(value).lower()
    value = re.sub(r"[^a-z0-9 ]+", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def parse_money(value: str) -> Optional[float]:
    if not value:
        return None
    raw = str(value).replace("\u20b9", "₹")
    # A dash means zero in these invoices: Freight -, IGST -, etc.
    if re.search(r"[:|]\s*-\s*$", raw) or raw.strip() in {"-", "—", "–"}:
        return 0.0
    matches = MONEY_RE.findall(raw)
    if not matches:
        return None
    num = matches[-1].replace(",", "")
    try:
        return round(float(num), 2)
    except ValueError:
        return None


def parse_date(value: str) -> str:
    if not value:
        return ""
    m = DATE_RE.search(value)
    candidate = m.group(1) if m else value
    try:
        dt = date_parser.parse(candidate, dayfirst=True, fuzzy=True)
        return dt.strftime("%d-%m-%Y")
    except Exception:
        return ""


def money(v: Optional[float]) -> float:
    return round(float(v or 0), 2)


@dataclass
class InvoiceRecord:
    record_type: str = "Sales"
    invoice_no: str = ""
    invoice_date: str = ""
    party_name: str = ""
    gstin: str = ""
    hsn: str = ""
    taxable_value: Optional[float] = None
    freight: float = 0.0
    cgst: Optional[float] = None
    sgst: Optional[float] = None
    igst: Optional[float] = None
    total: Optional[float] = None
    source_file: str = ""
    status: str = "OK"
    notes: str = ""
    confidence: str = "HIGH"

    def as_row(self) -> Dict[str, object]:
        data = asdict(self)
        for key in ["taxable_value", "freight", "cgst", "sgst", "igst", "total"]:
            if data[key] is None:
                data[key] = 0.0
        return data


@dataclass
class ItemRecord:
    source_file: str = ""
    invoice_no: str = ""
    description: str = ""
    hsn_sac: str = ""
    quantity: str = ""
    rate: str = ""
    taxable_value: Optional[float] = None
    tax_rate: str = ""
    total: Optional[float] = None

    def as_row(self) -> Dict[str, object]:
        data = asdict(self)
        for key in ["taxable_value", "total"]:
            if data[key] is None:
                data[key] = 0.0
        return data


class GSTInvoiceExtractor:
    def __init__(self, config: Dict):
        self.config = config or {}
        self.seller_gstin = str(self.config.get("business_gstin", SELLER_GSTIN_DEFAULT)).upper()

    def extract_file(self, path: Path, record_type: str = "Sales") -> Tuple[InvoiceRecord, List[ItemRecord]]:
        path = Path(path)
        text_lines, tables = self._read_file(path)
        rec, items = self._extract_demo_format(path, text_lines, tables, record_type)
        if rec is None:
            rec, items = self._extract_generic(path, text_lines, tables, record_type)
        self._final_validate(rec)
        return rec, items

    def _extract_demo_format(self, path: Path, text_lines: List[str], tables: List[List[List[str]]], record_type: str) -> Tuple[Optional[InvoiceRecord], List[ItemRecord]]:
        all_text = "\n".join(text_lines)
        if "bhilwara textile research center" not in all_text.lower():
            return None, []

        rec = InvoiceRecord(record_type=record_type, source_file=str(path))
        items: List[ItemRecord] = []

        # High-confidence direct regex extraction from full text.
        m = re.search(r"GST\s*Invoice\s*No\.?\s*[:\-]?\s*([A-Z0-9\-/]+)", all_text, re.I)
        if m:
            rec.invoice_no = clean_text(m.group(1))
        m = re.search(r"Invoice\s*Date\s*:\s*([^\n|]+)", all_text, re.I)
        if m:
            rec.invoice_date = parse_date(m.group(1))
        m = re.search(r"Name\s*:\s*(.*?)(?:Address\s*:|Mob\.\s*No|Phone\s*No|GSTIN\s*:)", all_text, re.I | re.S)
        if m:
            rec.party_name = clean_text(m.group(1)).strip(" ,")
        gstins = [g.upper() for g in GSTIN_RE.findall(all_text)]
        for g in gstins:
            if g != self.seller_gstin:
                rec.gstin = g
                break
        if not rec.gstin and gstins:
            rec.gstin = gstins[-1]

        # Parse cells in all tables. More reliable than paragraphs for Word invoices.
        for table in tables:
            for row in table:
                cells = [clean_text(c) for c in row]
                joined = " | ".join(cells)
                low = joined.lower()
                last_amount = parse_money(cells[-1] if cells else joined)

                if "total amt. before tax" in low or "total amt before tax" in low:
                    rec.taxable_value = last_amount
                elif "freight" in low:
                    rec.freight += money(last_amount)
                elif "add: sgst" in low or "sgst 9" in low or re.search(r"\bsgst\b", low):
                    rec.sgst = money(last_amount)
                elif "add: cgst" in low or "cgst 9" in low or re.search(r"\bcgst\b", low):
                    rec.cgst = money(last_amount)
                elif "add: igst" in low or re.search(r"\bigst\b", low):
                    rec.igst = money(last_amount)
                elif "total amt. after tax" in low or "total amt after tax" in low:
                    rec.total = last_amount

                # Item row detection: S.N. | Test Parameters | HSN Code | Quantity | Rs./sample | Amount.
                if len(cells) >= 6 and re.fullmatch(r"\d+", cells[0] or "") and re.search(r"\d{4,8}", cells[2] or ""):
                    item = ItemRecord(
                        source_file=str(path),
                        invoice_no=rec.invoice_no,
                        description=cells[1],
                        hsn_sac=re.search(r"\d{4,8}", cells[2]).group(0),
                        quantity=cells[3],
                        rate=cells[4],
                        taxable_value=parse_money(cells[5]),
                        tax_rate="18%",
                        total=parse_money(cells[5]),
                    )
                    items.append(item)
                    if not rec.hsn:
                        rec.hsn = item.hsn_sac

        # Fallback amount extraction from text, including labels where table extraction fails.
        label_patterns = {
            "taxable_value": [r"Total\s*Amt\.?\s*Before\s*Tax\s*([0-9,.]+)"],
            "freight": [r"Add\s*:\s*Freight\s*([0-9,.\-]+)"],
            "sgst": [r"Add\s*:\s*SGST\s*9%\s*([0-9,.\-]+)"],
            "cgst": [r"Add\s*:\s*CGST\s*9%\s*([0-9,.\-]+)"],
            "igst": [r"Add\s*:\s*IGST\s*([0-9,.\-]+)"],
            "total": [r"Total\s*Amt\.?\s*After\s*Tax\s*([0-9,.]+)"],
        }
        for field, pats in label_patterns.items():
            if field != "freight" and getattr(rec, field) not in (None, "", 0):
                continue
            for pat in pats:
                mm = re.search(pat, all_text, re.I)
                if mm:
                    val = parse_money(mm.group(1))
                    if field == "freight":
                        if rec.freight == 0:
                            rec.freight = money(val)
                    else:
                        setattr(rec, field, val)
                    break

        if rec.taxable_value is not None and rec.freight:
            rec.taxable_value = money(rec.taxable_value + rec.freight)

        for field in ("cgst", "sgst", "igst"):
            if getattr(rec, field) is None:
                setattr(rec, field, 0.0)

        if not rec.hsn:
            hsn_match = re.search(r"\b(99\d{4}|\d{6})\b", all_text)
            if hsn_match:
                rec.hsn = hsn_match.group(1)
            else:
                rec.hsn = "998346"

        return rec, items

    def _extract_generic(self, path: Path, text_lines: List[str], tables: List[List[List[str]]], record_type: str) -> Tuple[InvoiceRecord, List[ItemRecord]]:
        flat_text = "\n".join(text_lines)
        kv = self._extract_key_values(text_lines, tables)
        rec = InvoiceRecord(record_type=record_type, source_file=str(path))
        rec.invoice_no = self._find_by_keywords(kv, self.config.get("invoice_no_keywords", [])) or self._regex_invoice_no(flat_text)
        rec.invoice_date = self._find_date(kv, flat_text)
        rec.party_name = self._find_by_keywords(kv, self.config.get("party_keywords", [])) or self._guess_party(text_lines)
        rec.gstin = self._find_gstin(kv, flat_text)
        amount_keywords = self.config.get("amount_keywords", {})
        for field in ["taxable_value", "cgst", "sgst", "igst", "total"]:
            raw = self._find_by_keywords(kv, amount_keywords.get(field, []))
            setattr(rec, field, parse_money(raw or ""))
        self._fill_amounts_from_lines(rec, text_lines)
        items = self._extract_items(tables, path, rec.invoice_no)
        if items and not rec.hsn:
            rec.hsn = items[0].hsn_sac
        return rec, items

    def _read_file(self, path: Path) -> Tuple[List[str], List[List[List[str]]]]:
        suffix = path.suffix.lower()
        if suffix == ".docx":
            return self._read_docx(path)
        if suffix == ".pdf":
            return self._read_pdf(path)
        raise ValueError(f"Unsupported file type: {suffix}. Use .docx or .pdf")

    def _read_docx(self, path: Path) -> Tuple[List[str], List[List[List[str]]]]:
        doc = Document(str(path))
        lines = [clean_text(p.text) for p in doc.paragraphs if clean_text(p.text)]
        tables: List[List[List[str]]] = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([clean_text(cell.text) for cell in row.cells])
            if rows:
                tables.append(rows)
                for row in rows:
                    joined = " | ".join([c for c in row if c])
                    if joined:
                        lines.append(joined)
        return lines, tables

    def _read_pdf(self, path: Path) -> Tuple[List[str], List[List[List[str]]]]:
        if pdfplumber is None:
            raise RuntimeError("PDF support needs pdfplumber. Run: pip install pdfplumber")
        lines: List[str] = []
        tables: List[List[List[str]]] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                lines.extend([clean_text(x) for x in text.splitlines() if clean_text(x)])
                for table in page.extract_tables() or []:
                    rows = [[clean_text(c or "") for c in row] for row in table]
                    if rows:
                        tables.append(rows)
        return lines, tables

    def _extract_key_values(self, lines: List[str], tables: List[List[List[str]]]) -> Dict[str, str]:
        kv: Dict[str, str] = {}
        for line in lines:
            if ":" in line:
                left, right = line.split(":", 1)
                if len(left) <= 45 and right.strip():
                    kv[norm_key(left)] = clean_text(right)
        for table in tables:
            for row in table:
                cells = [c for c in row if c]
                if len(cells) >= 2:
                    for i in range(len(cells) - 1):
                        key = norm_key(cells[i])
                        if 2 <= len(key) <= 45 and cells[i + 1]:
                            kv[key] = clean_text(cells[i + 1])
        return kv

    def _find_by_keywords(self, kv: Dict[str, str], keywords: List[str]) -> str:
        for key, val in kv.items():
            for kw in keywords:
                if norm_key(kw) in key:
                    return clean_text(val)
        return ""

    def _regex_invoice_no(self, text: str) -> str:
        patterns = [
            r"(?:GST\s*)?(?:invoice|inv|bill)\s*(?:no|number|#)?\.?\s*[:\-]?\s*([A-Z0-9\-/]+)",
            r"\bNo\.?\s*[:\-]?\s*([A-Z0-9\-/]{2,})\b",
        ]
        for pat in patterns:
            m = re.search(pat, text, re.I)
            if m:
                return clean_text(m.group(1))
        return ""

    def _find_date(self, kv: Dict[str, str], text: str) -> str:
        raw = self._find_by_keywords(kv, self.config.get("date_keywords", []))
        dt = parse_date(raw)
        if dt:
            return dt
        m = DATE_RE.search(text)
        return parse_date(m.group(1)) if m else ""

    def _find_gstin(self, kv: Dict[str, str], text: str) -> str:
        raw = self._find_by_keywords(kv, self.config.get("gstin_keywords", []))
        gstins = [g.upper() for g in GSTIN_RE.findall(raw or "") + GSTIN_RE.findall(text or "")]
        for g in gstins:
            if g != self.seller_gstin:
                return g
        return gstins[-1] if gstins else clean_text(raw).upper()

    def _guess_party(self, lines: List[str]) -> str:
        bad_words = {"tax invoice", "invoice", "gstin", "date", "total", "amount", "bhilwara textile research center"}
        for line in lines[:25]:
            n = norm_key(line)
            if 4 < len(line) < 90 and not any(w in n for w in bad_words):
                return line
        return ""

    def _fill_amounts_from_lines(self, rec: InvoiceRecord, lines: List[str]) -> None:
        label_map = self.config.get("amount_keywords", {})
        for line in lines:
            n = norm_key(line)
            for field, keywords in label_map.items():
                if getattr(rec, field) in (None, 0) and any(norm_key(k) in n for k in keywords):
                    amount = parse_money(line)
                    if amount is not None:
                        setattr(rec, field, amount)
            if "freight" in n:
                amt = parse_money(line)
                if amt:
                    rec.freight += amt

    def _extract_items(self, tables: List[List[List[str]]], path: Path, invoice_no: str) -> List[ItemRecord]:
        items: List[ItemRecord] = []
        for table in tables:
            if len(table) < 2:
                continue
            header_idx = -1
            headers: List[str] = []
            for idx, row in enumerate(table[:6]):
                normalized = [norm_key(c) for c in row]
                joined = " ".join(normalized)
                if ("hsn" in joined or "sac" in joined) and ("amount" in joined or "rate" in joined or "quantity" in joined):
                    header_idx = idx
                    headers = normalized
                    break
            if header_idx == -1:
                continue
            for row in table[header_idx + 1:]:
                if not any(row):
                    continue
                joined = norm_key(" ".join(row))
                if any(x in joined for x in ["total", "grand total", "taxable value", "cgst", "sgst", "igst", "discount"]):
                    continue
                item = ItemRecord(source_file=str(path), invoice_no=invoice_no)
                for i, cell in enumerate(row):
                    head = headers[i] if i < len(headers) else ""
                    if any(x in head for x in ["description", "particular", "item", "goods", "service", "test parameters"]):
                        item.description = cell
                    elif "hsn" in head or "sac" in head:
                        m = re.search(r"\d{4,8}", cell or "")
                        item.hsn_sac = m.group(0) if m else cell
                    elif "qty" in head or "quantity" in head:
                        item.quantity = cell
                    elif "rs sample" in head or head == "rate" or " rate" in head:
                        item.rate = cell
                    elif "taxable" in head or "value" in head or "amount" in head:
                        item.taxable_value = parse_money(cell)
                    elif "total" in head:
                        item.total = parse_money(cell)
                    elif "tax" in head and "rate" in head:
                        item.tax_rate = cell
                if item.description or item.hsn_sac or item.taxable_value:
                    items.append(item)
        return items

    def _final_validate(self, rec: InvoiceRecord) -> None:
        notes: List[str] = []
        # Normalize money fields.
        for f in ("taxable_value", "cgst", "sgst", "igst", "total"):
            v = getattr(rec, f)
            if v is not None:
                setattr(rec, f, money(v))
        if rec.cgst is None:
            rec.cgst = 0.0
        if rec.sgst is None:
            rec.sgst = 0.0
        if rec.igst is None:
            rec.igst = 0.0

        # GST mode cleanup. In one normal invoice GST should be either:
        # - same state: CGST + SGST
        # - other state: IGST
        # Some Word tables can accidentally read the text/rate around "IGST -" as 18.00.
        seller_state = (self.seller_gstin or "")[:2]
        buyer_state = (rec.gstin or "")[:2]
        mixed_gst = bool(rec.igst and (rec.cgst or rec.sgst))
        if mixed_gst and seller_state and buyer_state:
            if buyer_state == seller_state:
                notes.append(f"Auto-fixed: ignored IGST {rec.igst} because buyer GSTIN is same-state and CGST/SGST are present")
                rec.igst = 0.0
            else:
                notes.append("Auto-fixed: ignored CGST/SGST because buyer GSTIN is inter-state and IGST is present")
                rec.cgst = 0.0
                rec.sgst = 0.0
        elif mixed_gst:
            notes.append("Mixed IGST with CGST/SGST found; please verify GST mode")

        # If taxable was missed or freight was not included, use official total - taxes.
        tax_sum = money(rec.cgst + rec.sgst + rec.igst)
        if rec.total not in (None, 0) and tax_sum >= 0:
            taxable_from_total = money((rec.total or 0) - tax_sum)
            if rec.taxable_value in (None, 0) and taxable_from_total > 0:
                rec.taxable_value = taxable_from_total
                notes.append("Taxable calculated from total minus GST")
            elif rec.taxable_value is not None:
                expected = money(rec.taxable_value + tax_sum)
                if abs(expected - money(rec.total)) > 1.0:
                    # Auto-correct when official total and tax values clearly imply another taxable value.
                    if taxable_from_total > 0:
                        notes.append(f"Taxable corrected from {rec.taxable_value} to {taxable_from_total} using total-GST check")
                        rec.taxable_value = taxable_from_total
                    else:
                        notes.append(f"Total mismatch: taxable+GST={expected}, invoice total={rec.total}")

        # If total missing, calculate it.
        if rec.total in (None, 0) and rec.taxable_value not in (None, 0):
            rec.total = money((rec.taxable_value or 0) + tax_sum)
            notes.append("Total calculated from taxable plus GST")

        # GST math check: Demo is generally 18%; allow rounding.
        if rec.taxable_value and (rec.cgst or rec.sgst) and not rec.igst:
            expected_half = money(rec.taxable_value * 0.09)
            if abs(expected_half - money(rec.cgst)) > 1.5 or abs(expected_half - money(rec.sgst)) > 1.5:
                notes.append("CGST/SGST does not match 9% each")
        if rec.taxable_value and rec.igst and not (rec.cgst or rec.sgst):
            expected_igst = money(rec.taxable_value * 0.18)
            if abs(expected_igst - money(rec.igst)) > 1.5:
                notes.append("IGST does not match 18%")

        # Required fields.
        missing = []
        for label, attr in [("Bill No", "invoice_no"), ("Date", "invoice_date"), ("Name", "party_name"), ("GSTIN", "gstin"), ("Taxable", "taxable_value"), ("Total", "total")]:
            val = getattr(rec, attr)
            if val in (None, "", 0, 0.0):
                missing.append(label)
        if missing:
            notes.insert(0, "Missing/uncertain: " + ", ".join(missing))

        # GSTIN format check.
        if rec.gstin and not GSTIN_RE.fullmatch(rec.gstin):
            notes.append("GSTIN format looks invalid")

        rec.notes = "; ".join(dict.fromkeys([n for n in notes if n]))
        # Auto-corrections are not errors. Mark CHECK only for real missing/mismatch notes.
        hard_notes = [n for n in notes if n.startswith("Missing") or "mismatch" in n.lower() or "invalid" in n.lower() or "does not match" in n.lower()]
        rec.status = "CHECK" if hard_notes else "OK"
        rec.confidence = "LOW" if rec.status == "CHECK" else "HIGH"


def _bill_sort_key(invoice_no: str):
    inv = str(invoice_no or "").strip()
    nums = re.findall(r"\d+", inv)
    if nums:
        return (0, int(nums[-1]), inv)
    return (1, 10**12, inv)


def scan_folder(input_folder: str | Path, record_type: str, config: Dict) -> Tuple[List[InvoiceRecord], List[ItemRecord]]:
    extractor = GSTInvoiceExtractor(config)
    folder = Path(input_folder)
    files = sorted(list(folder.rglob("*.docx")) + list(folder.rglob("*.pdf")), key=lambda p: _bill_sort_key(p.stem))
    records: List[InvoiceRecord] = []
    items: List[ItemRecord] = []
    for file in files:
        if file.name.startswith("~$"):
            continue
        try:
            rec, item_rows = extractor.extract_file(file, record_type=record_type)
            records.append(rec)
            items.extend(item_rows)
        except Exception as exc:
            records.append(InvoiceRecord(record_type=record_type, source_file=str(file), status="ERROR", notes=str(exc), confidence="LOW"))

    # Duplicate bill detection after extraction.
    seen: Dict[str, int] = {}
    for rec in records:
        key = str(rec.invoice_no).strip()
        if not key:
            continue
        seen[key] = seen.get(key, 0) + 1
    duplicates = {k for k, v in seen.items() if v > 1}
    for rec in records:
        if rec.invoice_no in duplicates:
            rec.status = "CHECK"
            rec.confidence = "LOW"
            rec.notes = (rec.notes + "; " if rec.notes else "") + "Duplicate bill number"

    records.sort(key=lambda r: _bill_sort_key(r.invoice_no))
    return records, items
